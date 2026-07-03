import os
import re
import tempfile
import requests
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def _download_image(url):
    if not url:
        return None
    try:
        r=requests.get(url,timeout=20)
        r.raise_for_status()
        f=tempfile.NamedTemporaryFile(delete=False,suffix=".jpg")
        f.write(r.content)
        f.close()
        return f.name
    except Exception:
        return None

def generate_report(report, hero, sections, filename):
    os.makedirs(os.path.dirname(filename), exist_ok=True)
    doc=Document()
    m=re.search(r"^#\s+(.+)$",report,re.MULTILINE)
    title=m.group(1).strip() if m else "Research Report"
    h=doc.add_heading(title,1)
    h.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    h.runs[0].font.size=Pt(22)
    h.runs[0].font.color.rgb=RGBColor(31,78,121)

    hp=_download_image(hero.get("image_url") if hero else None)
    if hp:
        doc.add_picture(hp,width=Inches(6.5))
        os.remove(hp)

    for section in sections:
        doc.add_heading(section["title"],2)
        img=section.get("image")
        ip=_download_image(img.get("image_url")) if img else None
        if ip:
            doc.add_picture(ip,width=Inches(5.8))
            os.remove(ip)
        for line in section["content"].splitlines():
            line=line.strip()
            if not line:
                continue
            line=re.sub(r"^#+\s*","",line)
            if line.startswith("- ") or line.startswith("* "):
                p=doc.add_paragraph(style="List Bullet")
                p.add_run(line[2:])
            else:
                p=doc.add_paragraph()
                parts=re.split(r"(\*\*.*?\*\*)",line)
                for part in parts:
                    if not part:
                        continue
                    if part.startswith("**") and part.endswith("**"):
                        r=p.add_run(part[2:-2]); r.bold=True
                    else:
                        p.add_run(part)
    doc.save(filename)
    return filename
